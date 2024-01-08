from plyer import notification

def send_notification(title, message, time):
    notification.notify(
            title=title,
            message=message,
            app_icon=None,
            timeout=time
        )
    
import webbrowser
import urllib.parse

def compose_email(to, subject, body):
    body = urllib.parse.quote(body)
    # Construct the mailto URL with parameters
    mailto_url = f"mailto:{to}?subject={subject}&body={body}"

    # Open the default web browser with the mailto URL
    webbrowser.open(mailto_url)


import datetime
from docx import Document
from docx.shared import Inches

def writeFile(text):

    # Get today's date in YYYY-MM-DD format
    todays_date = datetime.date.today().strftime("%Y-%m-%d")

    # Specify the folder path where you want to create the file
    folder_path = "E:/AI Boosted To Do List/" + todays_date  # Replace with your desired folder path

    # Create the folder if it doesn't exist
    import os
    os.makedirs(folder_path, exist_ok=True)

    # Create the document file name with today's date
    file_name = todays_date + ".docx"
    file_path = os.path.join(folder_path, file_name)

    # Create a new document
    document = Document()

    # Add a heading
    document.add_heading("Document for " + todays_date, 0)

    # Add a paragraph of text
    paragraph = document.add_paragraph(text)

    # Add some formatting to the paragraph
    paragraph.style = "List Bullet"  # Apply a bullet list style
    paragraph.paragraph_format.left_indent = Inches(0.5)  # Indent the paragraph

    # Add more paragraphs or content as needed

    # Save the document to the specified file path
    document.save(file_path)

    print(f"Document created successfully: {file_path}")
